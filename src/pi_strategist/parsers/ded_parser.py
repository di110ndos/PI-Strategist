"""Parser for DED documents (Word, Markdown, PDF)."""

import re
from pathlib import Path
from typing import Optional

from pi_strategist.models import (
    AcceptanceCriteria,
    DEDDocument,
    Epic,
    Story,
    Task,
)


class DEDParser:
    """Parser for Design & Engineering Documents."""

    # Patterns for extracting structure
    EPIC_PATTERNS = [
        r"(?:^|\n)#+\s*Epic[:\s]*(.+?)(?:\n|$)",
        r"(?:^|\n)Epic[:\s]+(.+?)(?:\n|$)",
        r"\[EPIC-(\d+)\][:\s]*(.+?)(?:\n|$)",
        r"EPIC-(\d+)[:\s]*(.+?)(?:\n|$)",
    ]

    STORY_PATTERNS = [
        r"(?:^|\n)#+\s*(?:User\s+)?Story[:\s]*(.+?)(?:\n|$)",
        r"(?:^|\n)(?:User\s+)?Story[:\s]+(.+?)(?:\n|$)",
        r"\[STORY-(\d+)\][:\s]*(.+?)(?:\n|$)",
        r"STORY-(\d+)[:\s]*(.+?)(?:\n|$)",
        r"US-(\d+)[:\s]*(.+?)(?:\n|$)",
    ]

    TASK_PATTERNS = [
        r"(?:^|\n)#+\s*Task[:\s]*(.+?)(?:\n|$)",
        r"\[TASK-(\d+)\][:\s]*(.+?)(?:\n|$)",
        r"TASK-(\d+)[:\s]*(.+?)(?:\n|$)",
    ]

    AC_PATTERNS = [
        r"(?:^|\n)(?:AC|Acceptance\s+Criteria?)[:\s]*\n?((?:[-*•]\s*.+?\n?)+)",
        r"(?:^|\n)Given\s+.+?\s+When\s+.+?\s+Then\s+.+?(?:\n|$)",
        r"(?:^|\n)[-*•]\s*(?:Given|When|Then|And)\s+.+?(?:\n|$)",
    ]

    def __init__(self):
        """Initialize the DED parser."""
        self._docx_available = self._check_docx()
        self._pdf_available = self._check_pdf()

    def _check_docx(self) -> bool:
        """Check if python-docx is available."""
        try:
            import docx  # noqa: F401
            return True
        except ImportError:
            return False

    def _check_pdf(self) -> bool:
        """Check if pdfplumber is available."""
        try:
            import pdfplumber  # noqa: F401
            return True
        except ImportError:
            return False

    def parse(self, file_path: str | Path) -> DEDDocument:
        """Parse a DED document from file.

        Args:
            file_path: Path to the DED document (.docx, .md, or .pdf)

        Returns:
            Parsed DEDDocument
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        # Check format before existence for better error messages
        supported_formats = {".docx", ".md", ".markdown", ".txt", ".pdf"}
        if suffix not in supported_formats:
            raise ValueError(f"Unsupported file format: {suffix}")

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if suffix == ".docx":
            return self._parse_docx(path)
        elif suffix in (".md", ".markdown", ".txt"):
            return self._parse_markdown(path)
        elif suffix == ".pdf":
            return self._parse_pdf(path)

    def parse_text(self, text: str, filename: str = "unknown") -> DEDDocument:
        """Parse DED content from text.

        Args:
            text: Raw text content of the DED
            filename: Source filename for reference

        Returns:
            Parsed DEDDocument
        """
        doc = DEDDocument(filename=filename, raw_text=text)
        doc.epics = self._extract_structure(text)
        return doc

    def _parse_docx(self, path: Path) -> DEDDocument:
        """Parse a Word document."""
        if not self._docx_available:
            raise ImportError("python-docx is required to parse .docx files")

        import docx

        document = docx.Document(path)
        text_parts = []

        for para in document.paragraphs:
            text_parts.append(para.text)

        # Also extract from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)

        full_text = "\n".join(text_parts)
        return self.parse_text(full_text, path.name)

    def _parse_markdown(self, path: Path) -> DEDDocument:
        """Parse a Markdown document."""
        text = path.read_text(encoding="utf-8")
        return self.parse_text(text, path.name)

    def _parse_pdf(self, path: Path) -> DEDDocument:
        """Parse a PDF document."""
        if not self._pdf_available:
            raise ImportError("pdfplumber is required to parse .pdf files")

        import pdfplumber

        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        full_text = "\n".join(text_parts)
        return self.parse_text(full_text, path.name)

    def _extract_structure(self, text: str) -> list[Epic]:
        """Extract epic/story/task structure from text."""
        epics = []

        # Try to find epics
        epic_matches = self._find_epics(text)

        if epic_matches:
            for epic_id, epic_name, epic_content in epic_matches:
                epic = Epic(id=epic_id, name=epic_name, description="")
                epic.stories = self._extract_stories(epic_content, epic_id)
                epics.append(epic)
        else:
            # No explicit epics found, treat as single implicit epic
            epic = Epic(id="EPIC-001", name="Main Epic", description="")
            epic.stories = self._extract_stories(text, "EPIC-001")
            if epic.stories:
                epics.append(epic)

        return epics

    def _find_epics(self, text: str) -> list[tuple[str, str, str]]:
        """Find all epics in text and their content."""
        results = []
        epic_counter = 1

        for pattern in self.EPIC_PATTERNS:
            matches = list(re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE))
            if matches:
                for i, match in enumerate(matches):
                    groups = match.groups()
                    if len(groups) == 2:
                        epic_id = f"EPIC-{groups[0]}"
                        epic_name = groups[1].strip()
                    else:
                        epic_id = f"EPIC-{epic_counter:03d}"
                        epic_name = groups[0].strip()
                        epic_counter += 1

                    # Get content until next epic
                    start = match.end()
                    end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
                    content = text[start:end]

                    results.append((epic_id, epic_name, content))
                break

        return results

    def _extract_stories(self, text: str, epic_id: str) -> list[Story]:
        """Extract stories from text section."""
        stories = []
        story_counter = 1

        # Find story boundaries
        story_matches = []
        for pattern in self.STORY_PATTERNS:
            matches = list(re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE))
            if matches:
                story_matches = matches
                break

        if story_matches:
            for i, match in enumerate(story_matches):
                groups = match.groups()
                if len(groups) == 2:
                    story_id = f"STORY-{groups[0]}"
                    story_name = groups[1].strip()
                else:
                    story_id = f"STORY-{story_counter:03d}"
                    story_name = groups[0].strip()
                    story_counter += 1

                # Get content until next story
                start = match.end()
                end = story_matches[i + 1].start() if i + 1 < len(story_matches) else len(text)
                content = text[start:end]

                story = Story(
                    id=story_id,
                    name=story_name,
                    description=content[:200].strip(),
                    epic_id=epic_id,
                )
                story.acceptance_criteria = self._extract_acceptance_criteria(content, story_id, epic_id)
                story.tasks = self._extract_tasks(content, story_id, epic_id)
                stories.append(story)
        else:
            # No explicit stories, try to extract ACs directly
            acs = self._extract_acceptance_criteria(text, "STORY-001", epic_id)
            if acs:
                story = Story(
                    id="STORY-001",
                    name="Main Story",
                    description="",
                    epic_id=epic_id,
                    acceptance_criteria=acs,
                )
                stories.append(story)

        return stories

    def _extract_acceptance_criteria(
        self, text: str, story_id: str, epic_id: str
    ) -> list[AcceptanceCriteria]:
        """Extract acceptance criteria from text."""
        acs = []
        ac_counter = 1

        # Look for AC sections
        ac_section_pattern = r"(?:Acceptance\s+Criteria?|AC)[:\s]*\n?((?:[-*•\d.]\s*.+?\n?)+)"
        ac_matches = re.finditer(ac_section_pattern, text, re.IGNORECASE | re.MULTILINE)

        for match in ac_matches:
            content = match.group(1)
            # Split into individual criteria
            items = re.split(r"\n[-*•\d.]+\s*", content)
            for item in items:
                item = item.strip()
                if item and len(item) > 5:
                    acs.append(
                        AcceptanceCriteria(
                            id=f"AC-{ac_counter:03d}",
                            text=item,
                            story_id=story_id,
                            epic_id=epic_id,
                        )
                    )
                    ac_counter += 1

        # Also look for Given/When/Then format
        gwt_pattern = r"(Given\s+.+?\s+When\s+.+?\s+Then\s+.+?)(?=\n(?:Given|$)|\Z)"
        gwt_matches = re.finditer(gwt_pattern, text, re.IGNORECASE | re.DOTALL)

        for match in gwt_matches:
            ac_text = match.group(1).strip()
            if ac_text:
                acs.append(
                    AcceptanceCriteria(
                        id=f"AC-{ac_counter:03d}",
                        text=ac_text,
                        story_id=story_id,
                        epic_id=epic_id,
                    )
                )
                ac_counter += 1

        # Look for bullet points that might be ACs
        if not acs:
            bullet_pattern = r"[-*•]\s*(.+?)(?=\n[-*•]|\n\n|\Z)"
            bullet_matches = re.finditer(bullet_pattern, text, re.DOTALL)
            for match in bullet_matches:
                item = match.group(1).strip()
                # Filter out items that are likely not ACs
                if item and len(item) > 10 and not item.lower().startswith(("note:", "todo:")):
                    acs.append(
                        AcceptanceCriteria(
                            id=f"AC-{ac_counter:03d}",
                            text=item,
                            story_id=story_id,
                            epic_id=epic_id,
                        )
                    )
                    ac_counter += 1

        return acs

    def _extract_tasks(
        self, text: str, story_id: str, epic_id: str
    ) -> list[Task]:
        """Extract tasks from text."""
        tasks = []
        task_counter = 1

        for pattern in self.TASK_PATTERNS:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                groups = match.groups()
                if len(groups) == 2:
                    task_id = f"TASK-{groups[0]}"
                    task_name = groups[1].strip()
                else:
                    task_id = f"TASK-{task_counter:03d}"
                    task_name = groups[0].strip()
                    task_counter += 1

                # Try to extract hours from task name
                hours = self._extract_hours(task_name)

                tasks.append(
                    Task(
                        id=task_id,
                        name=task_name,
                        hours=hours,
                        sprint="",
                        story_id=story_id,
                        epic_id=epic_id,
                    )
                )

        return tasks

    def _extract_hours(self, text: str) -> float:
        """Extract hour estimate from text."""
        # Look for patterns like "8h", "8 hours", "(8h)", etc.
        patterns = [
            r"(\d+(?:\.\d+)?)\s*(?:h|hr|hrs|hour|hours)",
            r"\((\d+(?:\.\d+)?)\s*(?:h|hr|hrs|hour|hours)?\)",
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return float(match.group(1))
        return 0.0
